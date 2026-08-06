#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Update all project documents, adding v2.0 changelog with fix details.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

UPDATE_DATE = "2026-07-29"
VERSION = "v2.0"

UPDATES = {
    "需求分析说明书.docx": {
        "title": "需求分析说明书",
        "sections": [
            ("V2.0 功能需求补充", [
                "前端界面需求：系统需要提供基于 Vue 3 的 SPA 前端界面，支持管理员、教师、学生三种角色的完整交互操作，替代原有依赖 API 文档页面的方式。",
                "班级管理需求：管理员可通过 Web 界面进行班级的增删改查，支持 Excel 批量导入班级和学生信息。",
                "知识点管理需求：管理员可对知识点进行 CRUD 管理，支持将题目关联到知识点。",
                "题目管理需求：管理员和教师可通过前端界面新增/编辑/删除题目，支持客观题（单选、多选、判断）和主观题（问答），支持选项管理和查重检测。",
                "学生查看结果需求：学生提交作业后可查看每道题的得分和教师评语。",
                "仪表盘统计需求：管理员首页需展示班级数、学生数、教师数、题库数、作业数、知识点数等实时统计数据。",
                "教师批改需求：教师查看学生答案时需显示题目类型，支持对主观题逐题评分并填写评语。",
                "登录信息需求：登录后返回用户真实姓名，支持首次登录强制修改密码的提示。"
            ]),
            ("V2.0 非功能需求补充", [
                "测试需求：系统需包含 JUnit 5 单元测试，覆盖核心 Service 层业务逻辑，测试数量不少于 30 个。",
                "容器化部署需求：系统需支持 Docker 容器化部署，提供 Dockerfile 和 docker-compose.yml 编排文件。",
                "兼容性需求：系统需同时支持 H2 内存数据库（开发环境）和 MySQL 8.0（生产环境）。",
                "安全性需求：登录失败锁定（5次/30分钟）、BCrypt 密码加密、JWT 无状态认证、审计日志记录。",
                "密码安全需求：批量导入学生的默认密码使用随机强密码生成，而非学号后6位弱密码。"
            ]),
            ("V2.0 审查发现的缺失与补充", [
                "缺失-集成测试：当前仅有 Service 层单元测试，缺少 Controller 层 MockMvc 集成测试。建议后续补充。",
                "缺失-压力测试：课程演示了 JMeter 压测（100-4000并发），项目目前未包含 JMeter 测试计划。建议补充。",
                "缺失-CI/CD：未包含 Jenkins 或 GitHub Actions 自动化流水线配置。建议后续添加。",
                "缺失-WAR打包：课程讲解了WAR包外部Tomcat部署，项目目前仅支持JAR打包。如需多项目共用服务器，需补充。",
                "缺失-HTTPS：未配置SSL/TLS，所有通信为HTTP明文。生产环境建议通过Nginx反向代理添加HTTPS。"
            ])
        ]
    },
    "软件设计说明书.docx": {
        "title": "软件设计说明书",
        "sections": [
            ("V2.0 前端架构设计", [
                "前端采用 Vue 3 + Element Plus + Axios 技术栈，以 CDN 方式加载，单文件 SPA 架构（index.html，约900行）。",
                "前端包含 10 个功能组件：AdminDashboard（仪表盘）、AdminTeachers（教师管理）、AdminClasses（班级管理）、AdminStudents（学生管理）、AdminQuestions（题库管理）、AdminKnowledge（知识点管理）、TeacherHomeworks（作业管理）、TeacherAssign（布置作业）、TeacherGrading（批改作业）、StudentHomeworks（学生作业）。",
                "前端通过 Axios 拦截器自动添加 JWT Bearer Token，401 响应自动跳转登录页。API 通信层封装了 40 个后端接口调用。",
                "组件间通信：TeacherAssign 通过 emit('assigned') 通知父组件切换视图；TeacherHomeworks 通过 emit('viewGrading') 预留批改导航。"
            ]),
            ("V2.0 后端 API 扩展", [
                "AdminController 新增 18 个端点：班级 CRUD（4个）、学生管理（2个）、题目 CRUD（5个）、知识点 CRUD（3个）、教学班管理（2个）、仪表盘（1个修复）、Excel导入（1个修复）。",
                "TeacherController 新增 3 个端点：教师作业列表、作业详情、学生答案查看（含questionType填充）。",
                "StudentController 新增 1 个端点：提交结果查看。",
                "UserService 新增 2 个方法：importClasses（Excel 导入班级）、importStudentsFromExcel（Excel 导入学生）。"
            ]),
            ("V2.0 关键设计决策与修复", [
                "SubmissionAnswer 新增 transient 字段 questionType：教师批改界面需显示题目类型（选择题/判断题/问答题），该字段通过 TeacherController 在查询答案后从 Question 表填充，不存储在 submission_answer 表中。",
                "AuthController /me 端点扩展：原仅返回 userId+role，现扩展为返回 username、realName、pwdResetRequired，供前端显示用户名和强制密码修改提示。",
                "登录后两步获取用户信息：前端先调用 /login 获取 token，再调用 /me 获取真实姓名。若 /me 调用失败则回退使用登录用户名。",
                "学生默认密码策略变更：从「学号后6位」改为调用 generateRandomPassword() 生成10位随机强密码，避免因学号位数不足导致密码不符合复杂度要求。",
                "Docker 健康检查修复：eclipse-temurin:17-jre-alpine 镜像不含 wget，在 Dockerfile 中增加 apk add --no-cache wget。"
            ]),
            ("V2.0 数据库设计更新", [
                "MySQL 初始化脚本 init.sql 表名从 'user' 修正为 'sys_user'，避免与 MySQL 保留关键字冲突。",
                "所有外键引用更新为 REFERENCES sys_user(id)。",
                "新增 ON DUPLICATE KEY UPDATE 处理，避免重复初始化报错。"
            ]),
            ("V2.0 部署架构设计", [
                "Docker 多阶段构建：Stage 1 使用 Maven 3.9 + JDK 17 编译打包，Stage 2 使用 JRE 17 运行。",
                "Docker Compose 编排：app 服务 + MySQL 8.0 服务，网络隔离（homework-net），数据卷持久化（mysql_data）。",
                "环境变量注入：DB_PASSWORD、JWT_SECRET 通过环境变量配置，默认值仅用于开发。"
            ])
        ]
    },
    "软件实现说明书.docx": {
        "title": "软件实现说明书",
        "sections": [
            ("V2.0 前端实现", [
                "前端为单个 index.html 文件（约900行），包含完整的 Vue 3 应用，使用 Composition API。",
                "使用 Element Plus UI 组件库：el-table（表格）、el-form（表单）、el-dialog（弹窗）、el-tag（标签）、el-menu（菜单）、el-upload（文件上传）、el-tabs（标签页）、el-radio-group/el-checkbox-group（选项组）等。",
                "登录页面：渐变背景 + 卡片式表单，角色识别后自动跳转对应菜单。登录成功后通过 /api/auth/me 获取真实姓名显示在顶栏。",
                "管理员功能：7 个菜单项，覆盖仪表盘（6项真实统计）、教师管理（增删改查+重置密码）、班级管理（增删改查+学生查看+Excel导入）、学生管理（按班级筛选+重置密码）、题库管理（增删改查+选项+知识点关联+启禁）、知识点管理（增删改查）。",
                "教师功能：3 个菜单项——作业管理（列表+提交情况弹窗+答案查看+成绩下载）、布置作业（教学班选择+题目搜索勾选+截止时间）、批改作业（选作业+逐题评分+评语）。",
                "学生功能：1 个菜单项——作业列表（待完成/已完成标签页）+ 在线作答（单选/多选/判断/问答）+ 查看结果（得分+评语）。",
                "题型支持：单选题（el-radio-group）、多选题（el-checkbox-group，答案拼接为字符串）、判断题（el-radio 对/错）、问答题（el-input textarea）。"
            ]),
            ("V2.0 后端实现", [
                "AdminController：从原有 95 行扩展到 270 行，新增班级、学生、题目、知识点、教学班的完整 CRUD 端点。仪表盘通过各 DAO.selectCount() 实时查询数据库。",
                "TeacherController：从原有 82 行扩展到约160行，新增教师作业列表（按teacherId筛选）、作业详情（含题目列表）、学生答案查看（含questionType填充）。",
                "StudentController：从原有 42 行扩展到 85 行，新增提交结果查看端点（返回Submission+SubmissionAnswer列表含得分评语）。",
                "UserServiceImpl：新增 importClasses 和 importStudentsFromExcel 方法，使用 Apache POI 解析 Excel 文件（支持字符串/数字/布尔类型单元格）。学生导入默认密码改为 generateRandomPassword() 随机强密码。",
                "AuthController /me 端点扩展：返回 userId、role、username、realName、pwdResetRequired。changePassword 增加 try/catch 返回结构化错误。",
                "SubmissionAnswer 新增 @TableField(exist=false) questionType 字段：在 TeacherController.getSubmissionAnswers 中通过 QuestionDao 查询填充。"
            ]),
            ("V2.0 关键代码修复记录", [
                "【严重】questionType 缺失修复：SubmissionAnswer 实体新增 transient 字段 questionType（@TableField(exist=false)），TeacherController.getSubmissionAnswers() 中遍历答案并通过 questionDao.selectById() 填充题目类型。修复前教师批改界面无法区分题型，Essay 评分功能完全失效。",
                "【高】前端静默错误修复：AdminDashboard 数据加载失败时添加 console.error 日志，避免统计数据显示全零而无任何提示。",
                "【高】AdminTeachers.del() 缺少 .catch() 修复：添加 .catch(()=>{}) 防止取消删除确认框时产生未处理的 Promise rejection。",
                "【中】changePassword 异常处理修复：原 RuntimeException 直接抛出导致 Spring 返回 500 错误，修复后通过 try/catch 捕获并返回 R.badRequest(e.getMessage())。",
                "【中】前端登录优化：登录成功后额外调用 /api/auth/me 获取真实姓名，原实现将用户名设为显示名。若 /me 调用失败则回退使用登录用户名。",
                "【中】api.submissionAnswers 修复：函数签名从 subId 单参数改为 (hwId, subId) 双参数，消除硬编码的 homeworkId=0。",
                "【中】Dockerfile wget 修复：eclipse-temurin:17-jre-alpine 不含 wget，添加 RUN apk add --no-cache wget 确保健康检查正常工作。",
                "【高】Windows 中文乱码修复：新建 DataInitializer 组件，使用 JDBC 参数化查询（PreparedStatement）以 Java 字符串直接插入中文数据，替代 Spring 的 sql.init 机制。根源是 Windows 平台 JVM 默认 GBK 编码导致 Spring ResourceDatabasePopulator 以错误编码读取 UTF-8 SQL 文件。Java 源码中的中文字符串由 javac 以 UTF-8 编译，通过 JDBC 传入 H2 始终保持 Unicode 正确。",
                "题目随机排序：HomeworkServiceImpl.getHomeworkDetail() 使用 studentId + homeworkId 作为随机种子，确保不同学生看到不同题目顺序但同一学生每次看到相同顺序。",
                "答题异常检测：SubmissionServiceImpl.submit() 如果答题时间 < 每题10秒，标记 suspiciousFlag = true，供教师人工审查。"
            ])
        ]
    },
    "软件测试说明书.docx": {
        "title": "软件测试说明书",
        "sections": [
            ("V2.0 单元测试详情", [
                "测试框架：JUnit 5 + Spring Boot Test + H2 内存数据库（jdbc:h2:mem:test）。",
                "测试配置：src/test/resources/application.yml 独立配置，使用与开发数据库同结构的 H2 测试数据库。",
                "测试数据：复用 src/main/resources/sql/data-h2.sql 中的初始化数据（1管理员+1教师+6学生+2班级+5题目+3知识点）。",
                "测试依赖：pom.xml 新增 spring-boot-starter-test、spring-security-test、h2（test scope）。",
                "",
                "UserServiceImplTest（10个用例）：",
                "  testLoginSuccess - 验证admin账号可正常登录并获取JWT Token",
                "  testLoginWrongPassword - 验证错误密码抛出异常",
                "  testLoginNonExistentUser - 验证不存在的用户抛出异常",
                "  testChangePasswordWeak - 验证弱密码（'123'）被拒绝",
                "  testResetPassword - 验证密码重置返回10位随机密码",
                "  testAddTeacher - 验证新增教师自动设置ROLE=TEACHER, STATUS=ACTIVE",
                "  testListTeachers - 验证教师列表查询不抛异常",
                "  testListTeachersWithKeyword - 验证关键字搜索功能",
                "  testImportStudents - 验证批量导入学生",
                "  testTransferStudent - 验证学生转班功能",
                "",
                "HomeworkServiceImplTest（8个用例）：",
                "  testAssignHomework - 验证作业布置（创建Homework+关联Question+初始化Submission）",
                "  testListStudentHomeworks - 验证学生作业列表查询",
                "  testGetHomeworkDetail - 验证作业详情含随机排序的题目和选项",
                "  testGetSubmissionStatus - 验证提交状态列表",
                "  testExportGrades - 验证成绩Excel导出（含表头和数据行）",
                "  testUpdateHomework - 验证作业修改",
                "  testDeleteHomework - 验证未提交作业可删除",
                "  testDeleteHomeworkWithSubmissions - 验证新布置作业删除",
                "",
                "QuestionServiceImplTest（10个用例）：",
                "  testAddSingleChoiceQuestion - 验证单选题新增（含选项）",
                "  testAddEssayQuestion - 验证问答题新增（含参考答案）",
                "  testAddQuestionWithoutAnswer - 验证客观题缺标准答案时抛出异常",
                "  testSearchQuestions - 验证题目搜索功能",
                "  testSearchByType - 验证按题型筛选",
                "  testCheckDuplicate - 验证查重方法不抛异常",
                "  testCheckNoDuplicate - 验证不相似题目查重返回空",
                "  testToggleStatus - 验证题目启用/禁用切换",
                "  testGetOptions - 验证选项查询（题目1有4个选项）",
                "  testBatchImport - 验证批量导入",
                "",
                "SubmissionServiceImplTest（4个用例）：",
                "  testSubmitWithCorrectAnswer - 验证提交正确答案自动评分满分",
                "  testSubmitWithWrongAnswer - 验证提交错误答案评分0分",
                "  testGetUngradedList - 验证未批改列表查询",
                "  testModifyAnswer - 验证修改答案后重新自动评分",
                "",
                "测试运行命令：mvn test",
                "测试结果：Tests run: 32, Failures: 0, Errors: 0, Skipped: 0 — BUILD SUCCESS"
            ]),
            ("V2.0 审查发现与待完善项", [
                "缺失-Controller层集成测试：当前所有测试为Service层单元测试，未使用MockMvc对HTTP端点进行集成测试。建议补充对登录、作业提交、批改等核心流程的端到端测试。",
                "缺失-压力测试：课程演示了JMeter压测（100-4000并发），项目目前未包含JMeter测试计划（.jmx文件）和压测报告。建议补充对核心API（如登录、作业提交）的压力测试。",
                "缺失-安全测试：未包含SQL注入测试、XSS测试、认证绕过测试、越权访问测试（学生访问教师接口等）。建议补充安全测试用例。",
                "缺失-回归测试套件：集成测试中的回归测试策略（每次替换桩/驱动后重新测试）未体现在代码中。",
                "注意-SubmissionAnswer.questionType字段：该字段为@TableField(exist=false)，不存储在数据库中，仅在TeacherController查询时动态填充。测试中不依赖此字段。",
                "注意-H2 CLOB字段限制：H2数据库中content字段类型为CLOB，LIKE查询可能存在限制，查重检测（checkDuplicate）的likeRight查询在H2环境下可能不完全准确。MySQL环境下正常工作。"
            ])
        ]
    },
    "系统使用说明书.docx": {
        "title": "系统使用说明书",
        "sections": [
            ("V2.0 前端使用说明", [
                "访问方式：启动后端后，浏览器打开 http://localhost:8080 即可访问前端界面。",
                "测试账号：admin（管理员） / T2024001（教师，张老师） / 20240001~20240006（学生），密码统一为 Admin123456。",
                "登录后顶部显示用户真实姓名（如'张老师'而非'T2024001'），角色标签颜色区分：红色=管理员、橙色=教师、绿色=学生。",
                "",
                "【管理员操作流程】",
                "1. 登录后进入仪表盘，查看系统统计数据（班级数、学生数、教师数、题库数、作业数、知识点数）",
                "2. 教师管理：新增/编辑/删除教师账号，重置密码（生成随机10位密码）",
                "3. 班级管理：新增/编辑/删除班级，点击'查看学生'进入学生列表，支持 Excel 导入学生",
                "4. 学生管理：选择班级查看学生列表，支持重置密码",
                "5. 题库管理：搜索/筛选题目，新增题目（选择题需配置选项和标准答案），删除/启禁题目",
                "6. 知识点管理：新增/编辑/删除知识点，支持名称、学科、描述字段",
                "",
                "【教师操作流程】",
                "1. 作业管理：查看已布置的作业列表，点击'提交情况'查看每位学生状态，点击'查看答案'批改，点击'下载成绩'导出Excel",
                "2. 布置作业：填写作业信息 -> 选择教学班级 -> 搜索勾选题目（支持按题型筛选）-> 设置截止时间 -> 点击'布置作业'",
                "3. 批改作业：选择作业 -> 查看学生列表 -> 点击'批改' -> 逐题输入得分和评语 -> 点击'保存评分'",
                "   注意：查看答案时每道题会显示题目类型标签（单选题/多选题/判断题/问答题），主观题可直接在弹窗中评分",
                "",
                "【学生操作流程】",
                "1. 我的作业：'待完成'标签显示未提交和可修改的作业，'已完成'标签显示已批改的作业",
                "2. 做作业：点击'去完成' -> 逐题作答 -> 点击'提交作业'。注意多选题需勾选多个选项",
                "3. 查看结果：在'已完成'标签找到已批改的作业 -> 点击'查看结果' -> 查看每道题的得分和教师评语"
            ]),
            ("V2.0 Docker 部署说明", [
                "前置条件：安装 Docker Desktop 并确保 Docker 服务运行中。",
                "国内用户建议配置镜像加速器（参考 xxq/daemon.json 中的镜像源列表）。",
                "部署步骤：",
                "  1. 在项目根目录执行：docker compose up -d",
                "  2. 首次启动会自动拉取 MySQL 8.0 镜像、编译应用、初始化数据库",
                "  3. 访问 http://localhost:8080",
                "  4. 查看日志：docker compose logs -f app",
                "  5. 停止服务：docker compose down",
                "  6. 如需清除数据：docker compose down -v",
                "环境变量（可选）：",
                "  DB_PASSWORD - MySQL root 密码，默认 Homework2026!",
                "  JWT_SECRET - JWT 签名密钥，生产环境务必更换"
            ]),
            ("V2.0 API 文档", [
                "Knife4j 接口文档地址：http://localhost:8080/doc.html",
                "H2 数据库控制台：http://localhost:8080/h2-console（JDBC URL: jdbc:h2:mem:homework，用户名sa，密码空）"
            ])
        ]
    },
    "队友环境配置指南.docx": {
        "title": "队友环境配置指南",
        "sections": [
            ("V2.0 新增配置方式", [
                "【方式一：本地运行（推荐开发使用）】",
                "1. JDK 17+（推荐 Eclipse Adoptium: https://adoptium.net）",
                "2. 项目内置 Maven 3.9.9（tools/maven/apache-maven-3.9.9/）",
                "3. 运行 start.bat（Windows）或执行：",
                "   tools/maven/apache-maven-3.9.9/bin/mvn spring-boot:run",
                "4. 访问 http://localhost:8080",
                "5. 运行测试：tools/maven/apache-maven-3.9.9/bin/mvn test",
                "6. 打包JAR：tools/maven/apache-maven-3.9.9/bin/mvn package -DskipTests",
                "",
                "【方式二：Docker 部署（推荐演示/生产使用）】",
                "1. 安装 Docker Desktop（https://www.docker.com/products/docker-desktop）",
                "2. 国内用户配置 Docker 镜像加速器（参考 xxq/daemon.json 中的镜像源列表，复制到 C:\\Users\\<用户名>\\.docker\\daemon.json）",
                "3. 在项目根目录执行：docker compose up -d",
                "4. 首次启动约需3-5分钟（拉取镜像+编译+初始化数据库）",
                "5. 访问 http://localhost:8080",
                "",
                "【前端开发环境（如需修改前端）】",
                "前端为单文件 HTML（src/main/resources/static/index.html），使用 CDN 加载 Vue 3 + Element Plus。",
                "只需文本编辑器即可修改，推荐 VS Code + Live Server 插件。",
                "如需离线开发，可将 CDN 资源下载到本地 static 目录并修改引用路径。",
                "注意：修改前端后无需重启后端，刷新浏览器即可看到更新。",
                "",
                "【数据库切换】",
                "开发环境默认使用 H2 内存数据库（无需安装，启动即用，数据不持久化）。",
                "生产环境需配置 MySQL 8.0，修改 src/main/resources/application.yml 中注释掉的 MySQL 数据源配置，",
                "或通过 Docker Compose 自动切换（app 容器会读取环境变量覆盖数据源配置）。",
                "MySQL 初始化脚本：src/main/resources/sql/init.sql（已使用 sys_user 表名，避免 MySQL 关键字冲突）。"
            ]),
            ("V2.0 已知注意事项（含 Windows 中文乱码解决方案）", [
                "Windows 中文乱码问题：Windows JVM 默认编码为 GBK，Spring Boot sql.init 以平台编码读取 UTF-8 SQL 文件导致中文损坏。项目已改用 Java DataInitializer（JDBC 参数化查询）替代 SQL 文件初始化，彻底解决乱码。",
                "如需修改测试数据：编辑 DataInitializer.java 中的 initData() 方法，中文字符串直接写在 Java 代码中（javac 默认 UTF-8 编译）。",
                "启动脚本配置：start.bat 已添加 MAVEN_OPTS=-Dfile.encoding=UTF-8 确保 JVM 以 UTF-8 运行。",
                "application.yml 已添加 server.servlet.encoding.force=true + charset=UTF-8 确保 HTTP 响应使用 UTF-8 编码。"
                "H2 数据库限制：题目查重功能（checkDuplicate）在 H2 中因 CLOB 字段 LIKE 查询限制可能不准确，MySQL 环境下正常工作。",
                "Docker 健康检查：Dockerfile 要求容器内 wget 可用（已通过 apk add 安装），如自定义基础镜像需注意。",
                "JWT 密钥：生产环境务必通过环境变量 JWT_SECRET 更换签名密钥，避免使用代码中的默认值。",
                "数据库密码：docker-compose.yml 中 MySQL root 密码默认值为 Homework2026!，生产环境务必通过 DB_PASSWORD 环境变量更换。",
                "端口冲突：应用默认 8080 端口，MySQL 映射到宿主机 3307 端口（避免与本地 MySQL 冲突）。如需修改，编辑 docker-compose.yml 中的 ports 配置。"
            ])
        ]
    }
}


def add_update_section(doc, sections):
    """Append v2.0 update record to the end of the document."""
    doc.add_page_break()

    heading = doc.add_heading(f'Appendix: {VERSION} Changelog', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Update Date: {UPDATE_DATE}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for section_title, items in sections:
        doc.add_heading(section_title, level=2)
        for item in items:
            if item == "":
                doc.add_paragraph()
            elif item.startswith("【"):
                doc.add_heading(item, level=3)
            else:
                p = doc.add_paragraph(item, style='List Bullet')


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))

    for filename, content in UPDATES.items():
        filepath = os.path.join(base_dir, filename)
        if not os.path.exists(filepath):
            print(f"[WARN] File not found, skip: {filename}")
            continue

        print(f"[UPDATE] {filename} ...")
        try:
            doc = Document(filepath)
            add_update_section(doc, content["sections"])

            backup = filepath + ".bak2"
            if not os.path.exists(backup):
                import shutil
                shutil.copy2(filepath, backup)
                doc.save(filepath)
                print(f"   [OK] Updated (backup saved as .bak2)")
            else:
                doc.save(filepath)
                print(f"   [OK] Updated (overwriting previous)")

        except Exception as e:
            print(f"   [ERROR] Update failed: {e}")

    print("\n[OK] All documents updated!")
    print("[INFO] See update_summary_v2.0.md for complete details.")


if __name__ == "__main__":
    main()
